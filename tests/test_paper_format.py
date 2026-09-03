import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


SCRIPTS = Path(__file__).resolve().parents[1] / "tools" / "docx" / "scripts"
sys.path.insert(0, str(SCRIPTS))

import paper_format as pf


class PaperFormatTests(unittest.TestCase):
    def _front_matter(self):
        doc = pf.new_document(contest="cumcm")
        pf.title(doc, "题目")
        pf.abstract_title(doc)
        pf.body(doc, "摘要正文")
        pf.keywords(doc, "优化；预测")
        return doc

    def test_reference_template_styles_are_kept_but_sample_body_is_cleared(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            source = Document()
            source.add_paragraph("模板示例正文，不应进入论文")
            source.save(template)

            doc = pf.new_document(contest="cumcm", template_path=template)

        self.assertNotIn("模板示例正文", "\n".join(p.text for p in doc.paragraphs))

    def test_official_fixed_template_content_can_be_preserved(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "official.docx"
            source = Document()
            source.add_paragraph("官方固定摘要页")
            source.save(template)

            doc = pf.new_document(
                contest="cumcm",
                template_path=template,
                preserve_template_content=True,
            )

        self.assertIn("官方固定摘要页", "\n".join(p.text for p in doc.paragraphs))

    def test_cumcm_structure_validator_requires_abstract_and_keywords(self):
        doc = pf.new_document(contest="cumcm")
        pf.title(doc, "题目")

        errors = pf.validate_paper_structure(doc, contest="cumcm")

        self.assertTrue(any("摘要" in error for error in errors))
        self.assertTrue(any("关键词" in error for error in errors))

    def test_complete_cumcm_front_matter_passes(self):
        doc = self._front_matter()

        errors = pf.validate_paper_structure(doc, contest="cumcm", quality_checks=False)

        self.assertEqual(errors, [])

    def test_quality_validation_reports_length_formula_figure_table_and_page_gaps(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(doc, contest="cumcm")

        for expected in ("9000", "公式", "图", "表", "渲染页数"):
            self.assertTrue(any(expected in issue for issue in issues), expected)
        self.assertTrue(any("低于质量目标 8" in issue for issue in issues))

    def test_other_contests_share_the_eight_figure_default(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            contest="mcm-icm",
            min_content_units=0,
            min_equations=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertTrue(any("低于质量目标 8" in issue for issue in issues))

    def test_table_caption_must_be_referenced_in_body(self):
        doc = self._front_matter()
        pf.body(doc, "正文没有引用下面的表格。")
        pf.three_line_table(doc, [["变量", "值"], ["x", "1"]])
        doc.add_paragraph("表1 参数结果")

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertTrue(any("表1" in issue and "正文" in issue for issue in issues))

    def test_reference_list_and_body_citations_are_bidirectionally_checked(self):
        doc = self._front_matter()
        pf.body(doc, "已有研究支持该方法[1]，但错误引用了[3]。")
        doc.add_paragraph("参考文献")
        doc.add_paragraph("[1] A. Author. A useful paper.")
        doc.add_paragraph("[2] B. Author. An uncited paper.")

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertTrue(any("[3]" in issue and "参考文献表" in issue for issue in issues))
        self.assertTrue(any("[2]" in issue and "未在正文引用" in issue for issue in issues))

    def test_compound_reference_citations_are_recognized(self):
        doc = self._front_matter()
        pf.body(doc, "相关方法见文献[1, 2]及文献[3-4]。")
        doc.add_paragraph("参考文献")
        for number in range(1, 5):
            doc.add_paragraph(f"[{number}] Reference {number}.")

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertFalse(any("未在正文引用" in issue for issue in issues))

    def test_rendered_page_limits_distinguish_target_from_official_maximum(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            rendered_pages=31,
        )

        self.assertTrue(any("官方上限" in issue and "30" in issue for issue in issues))

    def test_safe_save_rejects_skill_root(self):
        doc = self._front_matter()

        with self.assertRaisesRegex(ValueError, "PROJECT_ROOT"):
            pf.save_document(doc, pf.SKILL_ROOT, contest="cumcm")

    def test_completion_gate_rejects_incomplete_docx(self):
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "完整论文.docx"
            self._front_matter().save(path)

            report = pf.validate_document(path, contest="cumcm", rendered_pages=7)

        self.assertFalse(report["passed"])
        self.assertLess(report["metrics"]["content_units"], 9000)
        self.assertTrue(any("9000" in issue for issue in report["issues"]))

    def test_three_line_table_five_rules(self):
        from docx.oxml.ns import qn

        doc = pf.new_document(contest="cumcm")
        table = pf.three_line_table(doc, [["符号", "说明"], [[("math", r"x_i")], "决策变量"]])

        # 细三线：顶线 1pt（sz=8），表头分隔线 0.5pt（sz=4）
        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        self.assertEqual(borders.find(qn("w:top")).get(qn("w:sz")), "8")
        self.assertEqual(borders.find(qn("w:bottom")).get(qn("w:sz")), "8")
        # 首行跨页重复 + 全部行不跨页断裂
        first_tr_pr = table.rows[0]._tr.find(qn("w:trPr"))
        self.assertIsNotNone(first_tr_pr.find(qn("w:tblHeader")))
        for row in table.rows:
            self.assertIsNotNone(row._tr.find(qn("w:trPr")).find(qn("w:cantSplit")))
        # 单元格：垂直居中、单倍行距、五号
        for row in table.rows:
            for cell in row.cells:
                self.assertIsNotNone(
                    cell._tc.find(qn("w:tcPr")).find(qn("w:vAlign"))
                )
                for p in cell.paragraphs:
                    self.assertEqual(p.paragraph_format.line_spacing, 1.0)
        # 符号列混排公式：单元格内含行内 oMath
        cell_xml = table.cell(1, 0)._tc
        self.assertIsNotNone(cell_xml.find(f".//{qn('m:oMath')}"))

    def test_equation_display_centered_with_number(self):
        from docx.oxml.ns import qn
        from docx.enum.text import WD_TAB_ALIGNMENT

        doc = pf.new_document(contest="cumcm")
        pf.equation(doc, r"\min_{\theta} S(\theta)=\sum_{i=1}^{n} x_i", number="1")
        p = doc.paragraphs[-1]
        para = p._element
        # 双制表位方案：居中制表位（版心一半）+ 右制表位（版心右端），
        # 公式为行内 oMath（display 内容），编号悬挂右端。
        # 不用 oMathPara+jc=center——与编号同段时 Word 实测退化为左对齐。
        self.assertIsNone(para.find(qn("m:oMathPara")))
        stops = p.paragraph_format.tab_stops
        self.assertEqual(len(stops), 2)
        self.assertEqual(stops[0].alignment, WD_TAB_ALIGNMENT.CENTER)
        self.assertEqual(stops[0].position.twips, 4150)
        self.assertEqual(stops[1].alignment, WD_TAB_ALIGNMENT.RIGHT)
        self.assertEqual(stops[1].position.twips, 8300)
        self.assertIsNotNone(para.find(qn("m:oMath")))
        self.assertIsNotNone(para.find(f".//{qn('m:nary')}"))
        self.assertIsNotNone(para.find(f".//{qn('m:limLow')}"))
        self.assertIn("(1)", p.text)

    def test_code_block_is_paragraphs_not_table(self):
        from docx.oxml.ns import qn

        doc = pf.new_document(contest="cumcm")
        tables_before = len(doc.tables)
        pf.code_block(doc, "def f(x):\n    return x  # 注释\n")
        # 代码块用段落边框实现，不得新增表格（避免门禁误计为"表"）
        self.assertEqual(len(doc.tables), tables_before)
        texts = [p.text for p in doc.paragraphs if "def" in p.text]
        self.assertTrue(texts)
        run = doc.paragraphs[-2].runs[1]
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        self.assertEqual(fonts.get(qn("w:ascii")), "Consolas")
        # 关键字蓝色
        self.assertIsNotNone(doc.paragraphs[-2].runs[1].font.color.rgb)


if __name__ == "__main__":
    unittest.main()
